from __future__ import annotations

import copy
import os
import shutil
import zipfile
from pathlib import Path
import sys

sys.path.insert(0, r'C:\Users\Admin\.cache\codex-runtimes\codex-primary-runtime\dependencies\python')

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(r"D:\Sumo\sumo_train")
DOCX_IN = ROOT / "docs" / "dissertation" / "full_draft_submission_v7_clean.docx"
DOCX_OUT = ROOT / "docs" / "dissertation" / "full_draft_submission_v8.docx"
DOCX_TMP = ROOT / "docs" / "dissertation" / "_full_draft_submission_v8_work.docx"
FIG3_OUT = ROOT / "docs" / "dissertation" / "figures" / "final" / "figure_3_provider_success_fallback_v8.png"


def set_paragraph_text(paragraph, text: str) -> None:
    paragraph.text = text


def clear_paragraph(paragraph) -> None:
    paragraph.text = ""


def add_toc_field(paragraph, heading_text: str = "Table of Contents") -> None:
    p = paragraph._p
    for child in list(p):
        p.remove(child)

    run = paragraph.add_run(heading_text)
    run.add_break()

    field_run = paragraph.add_run()
    r = field_run._r

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and update field in Word."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    r.append(begin)
    r.append(instr)
    r.append(separate)
    r.append(placeholder)
    r.append(end)


def set_update_fields(doc: Document) -> None:
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def remove_last_column(table) -> None:
    tbl = table._tbl
    for tr in tbl.tr_lst:
        cells = list(tr.tc_lst)
        if cells:
            tr.remove(cells[-1])
    grid = tbl.tblGrid
    cols = list(grid.gridCol_lst)
    if cols:
        grid.remove(cols[-1])


def remove_last_rows(table, count: int) -> None:
    for _ in range(count):
        tbl = table._tbl
        rows = tbl.tr_lst
        if not rows:
            return
        tbl.remove(rows[-1])


def add_rows(table, total_rows: int) -> None:
    while len(table.rows) < total_rows:
        table.add_row()


def set_table_cell(table, row: int, col: int, text: str) -> None:
    table.rows[row].cells[col].text = text


def generate_figure_3(out_path: Path) -> None:
    categories = [
        ("Raw LLM 4V", 5.86, 94.14),
        ("Raw LLM 8V", 0.22, 99.78),
        ("Hybrid 4V", 4.95, 95.05),
        ("Hybrid 8V", 0.11, 99.89),
        ("Hybrid + Safety 4V", 4.95, 95.05),
        ("Hybrid + Safety 8V", 0.11, 99.89),
    ]

    width, height = 1800, 1050
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)

    def font(size, bold=False):
        candidates = [
            r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
            r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf",
        ]
        for candidate in candidates:
            try:
                return ImageFont.truetype(candidate, size)
            except OSError:
                continue
        return ImageFont.load_default()

    title_font = font(44, True)
    label_font = font(24, False)
    small_font = font(22, False)
    bold_small = font(22, True)

    draw.text((70, 40), "Provider success and fallback rate by LLM controller and scale", fill="black", font=title_font)
    draw.text((70, 100), "8V pooled total: 4 / 2784 = 0.14% success, 99.86% fallback", fill="#444444", font=label_font)

    left, right, top, bottom = 120, 1680, 200, 870
    chart_w = right - left
    chart_h = bottom - top
    draw.rectangle([left, top, right, bottom], outline="#333333", width=2)

    for pct in range(0, 101, 20):
        y = bottom - chart_h * pct / 100
        draw.line([left, y, right, y], fill="#E0E0E0", width=2)
        label = f"{pct}%"
        bbox = draw.textbbox((0, 0), label, font=small_font)
        draw.text((left - 20 - (bbox[2] - bbox[0]), y - 12), label, fill="#555555", font=small_font)

    bar_w = 145
    gap = 74
    group_gap = 48
    start_x = left + 45
    x = start_x
    success_color = (46, 125, 50)
    fallback_color = (249, 168, 37)
    legend_x = right - 420
    draw.rectangle([legend_x, 92, legend_x + 22, 114], fill=success_color)
    draw.text((legend_x + 34, 88), "Provider success", fill="black", font=label_font)
    draw.rectangle([legend_x, 126, legend_x + 22, 148], fill=fallback_color)
    draw.text((legend_x + 34, 122), "Fallback rate", fill="black", font=label_font)

    for idx, (label, success, fallback) in enumerate(categories):
        bar_left = x
        bar_height_success = chart_h * success / 100
        bar_height_fallback = chart_h * fallback / 100
        success_top = bottom - bar_height_success
        fallback_top = success_top - bar_height_fallback
        draw.rectangle([bar_left, success_top, bar_left + bar_w, bottom], fill=success_color)
        draw.rectangle([bar_left, fallback_top, bar_left + bar_w, success_top], fill=fallback_color)
        draw.rectangle([bar_left, top, bar_left + bar_w, bottom], outline="#333333", width=1)

        s_text = f"{success:.2f}%"
        f_text = f"{fallback:.2f}%"
        sb = draw.textbbox((0, 0), s_text, font=bold_small)
        fb = draw.textbbox((0, 0), f_text, font=bold_small)
        draw.text((bar_left + bar_w / 2 - (sb[2] - sb[0]) / 2, success_top + max(6, bar_height_success / 2 - 12)), s_text, fill="white", font=bold_small)
        draw.text((bar_left + bar_w / 2 - (fb[2] - fb[0]) / 2, fallback_top + max(6, bar_height_fallback / 2 - 12)), f_text, fill="#3E2723", font=bold_small)

        lb = draw.textbbox((0, 0), label, font=small_font)
        draw.text((bar_left + bar_w / 2 - (lb[2] - lb[0]) / 2, bottom + 22), label, fill="black", font=small_font)
        x += bar_w + gap
        if idx == 1 or idx == 3:
            x += group_gap

    draw.text((70, 930), "The figure excludes the rule-based baseline and focuses only on live LLM controllers.", fill="#444444", font=label_font)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    img.save(out_path)


def patch_media(docx_path: Path, media_name: str, replacement_path: Path) -> None:
    temp_path = docx_path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(docx_path, "r") as src, zipfile.ZipFile(temp_path, "w", compression=zipfile.ZIP_DEFLATED) as dst:
        for info in src.infolist():
            data = src.read(info.filename)
            if info.filename == media_name:
                data = replacement_path.read_bytes()
            dst.writestr(info, data)
    temp_path.replace(docx_path)


def main() -> None:
    generate_figure_3(FIG3_OUT)

    doc = Document(str(DOCX_IN))
    doc.core_properties.title = "A Structured LLM-Assisted Decision Pipeline for Unsignalised Intersection Control in SUMO"
    doc.core_properties.subject = "MSc Robotics dissertation"

    # Title page and TOC
    set_paragraph_text(
        doc.paragraphs[0],
        "University of Bristol MSc Robotics Dissertation\n"
        "Title: A Structured LLM-Assisted Decision Pipeline for Unsignalised Intersection Control in SUMO\n"
        "Author: [NEEDS_USER_CONFIRMATION]\n"
        "Supervisor: [NEEDS_USER_CONFIRMATION]\n"
        "Submission date: [NEEDS_USER_CONFIRMATION]",
    )
    add_toc_field(doc.paragraphs[1], "Table of Contents")

    # Chapter 3
    set_paragraph_text(
        doc.paragraphs[118],
        "For the LLM-assisted controllers, the structured prompt is submitted to a live language-model provider. "
        "The formal experiments used Groq with the openai/gpt-oss-20b model.",
    )
    set_paragraph_text(
        doc.paragraphs[119],
        "The request configuration was fixed throughout the evaluated experiments: provider Groq; model openai/gpt-oss-20b; "
        "max completion tokens 256; reasoning effort low; timeout 30 s; max retries 0.",
    )
    set_paragraph_text(
        doc.paragraphs[168],
        "The same frozen request settings were used for all live LLM calls so that controller differences reflected downstream "
        "processing rather than request drift. The configuration was Groq / openai/gpt-oss-20b / 256 max completion tokens / "
        "low reasoning effort / 30 s timeout / 0 retries.",
    )
    set_paragraph_text(
        doc.paragraphs[169],
        "Within the evaluated experiments, temperature, top_p, seed, and reasoning_format were not varied.",
    )

    # Chapter 4
    set_paragraph_text(doc.paragraphs[147], "4.2 Experimental Factors")
    set_paragraph_text(
        doc.paragraphs[148],
        "Three experimental factors were considered in the formal design: controller architecture, vehicle scale, and random seed.",
    )
    set_paragraph_text(
        doc.paragraphs[149],
        "The four controller architectures were Rule-based, Raw LLM, Hybrid LLM, and Hybrid LLM + Safety. The controller factor "
        "therefore compared a deterministic baseline with three increasingly staged LLM-assisted pipelines.",
    )
    set_paragraph_text(
        doc.paragraphs[150],
        "Each controller was evaluated at two vehicle scales, 4V and 8V, and with three random seeds, giving 4 × 2 × 3 = 24 experimental conditions.",
    )
    set_paragraph_text(doc.paragraphs[151], "4.3 Experimental Evidence and Run Validity")
    set_paragraph_text(
        doc.paragraphs[152],
        "The final dissertation evidence combines 12 valid 4V runs from the original formal_v2 batch with 12 corrected 8V runs from formal_v4, for 24 valid runs in total.",
    )
    set_paragraph_text(
        doc.paragraphs[153],
        "The nominal 8V traces from formal_v2 are excluded because trace validation showed that the underlying scenario had only four observed vehicles rather than eight.",
    )
    set_paragraph_text(
        doc.paragraphs[154],
        "The corrected 8V runs were executed after the configuration error was identified, while preserving the controller algorithms, prompt, request configuration, and decision semantics.",
    )
    set_paragraph_text(
        doc.paragraphs[155],
        "The final evidence used in this dissertation is therefore 12 valid 4V runs from formal_v2 and 12 corrected valid 8V runs from formal_v4.",
    )
    set_paragraph_text(
        doc.paragraphs[156],
        "The invalid nominal 8V formal_v2 runs and the intermediate rerun are excluded from the final comparative results.",
    )

    # Chapter 5
    set_paragraph_text(
        doc.paragraphs[211],
        "The results in this chapter are limited to the final dissertation evidence boundary and are reported as descriptive comparisons across controller and scale.",
    )
    set_paragraph_text(
        doc.paragraphs[212],
        "The usable dataset combines 12 valid 4V runs from formal_v2 with 12 corrected 8V runs from formal_v4.",
    )
    set_paragraph_text(
        doc.paragraphs[213],
        "The excluded nominal 8V formal_v2 traces and the intermediate rerun are retained only as validity history, not as evidence for the final comparison.",
    )
    set_paragraph_text(doc.paragraphs[214], "5.1 Experimental Evidence and Validity")
    set_paragraph_text(
        doc.paragraphs[215],
        "The final evidence set is deliberately narrow: it covers a single SUMO intersection, two vehicle scales, four controllers, and three seeds per condition.",
    )
    set_paragraph_text(
        doc.paragraphs[216],
        "This scope is sufficient for descriptive comparison, but not for inferential claims about universal model superiority or broader traffic regimes.",
    )
    set_paragraph_text(
        doc.paragraphs[217],
        "The results therefore describe the behaviour of the complete pipeline under a frozen experiment, not an all-purpose traffic benchmark.",
    )
    set_paragraph_text(doc.paragraphs[220], "5.2 Traffic Performance")
    set_paragraph_text(
        doc.paragraphs[221],
        "Table 2 summarises the traffic outcomes for the final evidence set. The table reports completion rate, mean waiting time, mean speed, throughput, and collision count using the aggregated seed means and standard deviations.",
    )
    set_paragraph_text(doc.paragraphs[222], "Table 2. Traffic performance by controller and scale")
    set_paragraph_text(
        doc.paragraphs[224],
        "Across both scales, the rule-based baseline shows markedly higher waiting time than the LLM-assisted controllers.",
    )
    set_paragraph_text(
        doc.paragraphs[225],
        "At 4V, the baseline records 82.000 steps of mean waiting time, whereas the LLM-assisted controllers record 15.000 steps.",
    )
    set_paragraph_text(
        doc.paragraphs[226],
        "At 8V, the baseline rises to 242.042 steps, while the LLM-assisted controllers remain at 15.292 steps.",
    )
    set_paragraph_text(
        doc.paragraphs[227],
        "Mean speed follows the same pattern, with the baseline dropping to 1.189 m/s at 8V while the LLM-assisted controllers remain at 6.599 m/s.",
    )
    set_paragraph_text(
        doc.paragraphs[228],
        "Completion rate is 100% in every valid cell and collision count remains zero throughout the final evidence.",
    )
    set_paragraph_text(doc.paragraphs[229], "5.3 Provider and Parser Reliability")
    set_paragraph_text(doc.paragraphs[230], "Table 3. Provider/parser/fallback reliability")
    set_paragraph_text(
        doc.paragraphs[232],
        "Provider success is low in every live LLM-bearing cell, and the corrected 8V boundary records only 4 successful provider calls out of 2,784 attempts across the three live-provider controllers.",
    )
    set_paragraph_text(
        doc.paragraphs[233],
        "Successful responses used finish_reason = stop and finite token usage, while parser success given provider success remained 100% in the corrected evidence.",
    )
    set_paragraph_text(
        doc.paragraphs[234],
        "The fallback path therefore dominates the live execution history, and provider failures map directly to fallback usage in the final evidence boundary.",
    )
    set_paragraph_text(
        doc.paragraphs[235],
        "This means that the traffic results should be interpreted as pipeline behaviour under constrained provider availability rather than as a clean measurement of model-only capability.",
    )
    clear_paragraph(doc.paragraphs[236])
    clear_paragraph(doc.paragraphs[237])
    set_paragraph_text(doc.paragraphs[238], "5.4 Decision-Flow Behaviour")
    set_paragraph_text(
        doc.paragraphs[239],
        "The trace schema preserves raw, validated, postprocessed, and final decisions separately, which makes it possible to describe the complete control pipeline rather than only the final action.",
    )
    set_paragraph_text(doc.paragraphs[240], "Table 4. Decision-source / postprocessor / safety behaviour")
    set_paragraph_text(
        doc.paragraphs[241],
        "The retained traces show that the final decision is usually inherited from fallback handling rather than from a live postprocessor or safety override.",
    )
    set_paragraph_text(doc.paragraphs[242], "5.5 Safety Outcomes")
    set_paragraph_text(doc.paragraphs[243], "Collision count is zero in every valid run.")
    set_paragraph_text(doc.paragraphs[244], "Safety override count is zero in every valid run.")
    set_paragraph_text(
        doc.paragraphs[245],
        "The safety verifier remained logged and traceable, but the corrected formal evidence does not show it being meaningfully exercised.",
    )
    set_paragraph_text(doc.paragraphs[246], "5.6 Summary")
    set_paragraph_text(
        doc.paragraphs[247],
        "The retained evidence combines a strong traffic advantage for the LLM-assisted pipeline with weak provider reliability and negligible downstream intervention by the cooperative and safety stages.",
    )
    set_paragraph_text(
        doc.paragraphs[248],
        "The identical 8V traffic results for Raw LLM, Hybrid, and Hybrid + Safety are consistent with fallback dominance and very low provider success.",
    )
    set_paragraph_text(
        doc.paragraphs[249],
        "Their similarity should be read as a consequence of the execution path, not as evidence that the controller variants are intrinsically equivalent.",
    )
    set_paragraph_text(
        doc.paragraphs[250],
        "Figures 1 to 4 summarise the corrected evidence in compact form.",
    )
    set_paragraph_text(
        doc.paragraphs[251],
        "Figure 1 illustrates mean waiting time across controller and scale.",
    )
    set_paragraph_text(
        doc.paragraphs[252],
        "Figure 2 illustrates mean speed across controller and scale.",
    )
    set_paragraph_text(
        doc.paragraphs[253],
        "Figure 3 illustrates provider success and fallback rate for the live LLM controllers only.",
    )
    set_paragraph_text(
        doc.paragraphs[254],
        "Figure 4 illustrates live-provider latency across the live LLM controllers.",
    )
    set_paragraph_text(
        doc.paragraphs[255],
        "These figures provide a compact visual summary of the corrected dissertation evidence.",
    )
    set_paragraph_text(
        doc.paragraphs[257],
        "Figure 1 shows that the waiting-time gap between the rule-based baseline and the LLM-assisted pipelines persists across both vehicle scales.",
    )
    set_paragraph_text(
        doc.paragraphs[260],
        "Figure 2 shows that the same pattern appears in mean speed, with the LLM-assisted pipelines maintaining substantially higher speed than the baseline.",
    )
    set_paragraph_text(
        doc.paragraphs[263],
        "Figure 3 is limited to the live LLM controllers and shows the extreme fallback dominance that underpins the reliability discussion in this chapter.",
    )
    set_paragraph_text(
        doc.paragraphs[266],
        "Figure 4 shows that the surviving live-provider calls are few but still measurable in latency, which is useful for understanding execution cost rather than control quality.",
    )

    # Chapter 6
    set_paragraph_text(
        doc.paragraphs[268],
        "6 Discussion",
    )
    set_paragraph_text(
        doc.paragraphs[269],
        "This chapter interprets the final dissertation evidence in relation to the research questions and the recovered literature.",
    )
    set_paragraph_text(
        doc.paragraphs[270],
        "The discussion stays bounded by the corrected evidence set: 12 valid 4V runs from formal_v2 and 12 corrected 8V runs from formal_v4.",
    )
    set_paragraph_text(doc.paragraphs[271], "6.1 Traffic Efficiency and RQ1")
    set_paragraph_text(
        doc.paragraphs[272],
        "RQ1 asks whether the LLM-assisted architecture improves traffic efficiency relative to rule-based control.",
    )
    set_paragraph_text(
        doc.paragraphs[273],
        "The corrected evidence answers that question cautiously in the affirmative at the pipeline level: the LLM-assisted controllers achieve lower waiting time and higher mean speed than the rule-based baseline in both tested scales.",
    )
    set_paragraph_text(doc.paragraphs[274], "4V rule-based: waiting 82.0 steps, speed 2.3098 m/s")
    set_paragraph_text(doc.paragraphs[275], "4V LLM-assisted: waiting 15.0 steps, speed 6.8026 m/s")
    set_paragraph_text(doc.paragraphs[276], "8V rule-based: waiting 242.0417 steps, speed 1.1895 m/s")
    set_paragraph_text(doc.paragraphs[277], "8V LLM-assisted: waiting 15.2917 steps, speed 6.5991 m/s")
    set_paragraph_text(
        doc.paragraphs[278],
        "The improvement should not be read as evidence that the LLM component alone caused the gain, because live-provider availability is too weak for that attribution.",
    )
    set_paragraph_text(
        doc.paragraphs[279],
        "The correct interpretation is therefore that the full pipeline behaved better than the rule-based baseline in the tested scenarios.",
    )
    set_paragraph_text(
        doc.paragraphs[280],
        "This is a traffic result about a pipeline, not a claim of intrinsic language-model superiority.",
    )
    set_paragraph_text(
        doc.paragraphs[281],
        "The fallback-heavy execution path remains the dominant explanatory factor.",
    )
    set_paragraph_text(doc.paragraphs[282], "6.2 Cooperative Post-processing and RQ2")
    set_paragraph_text(
        doc.paragraphs[283],
        "RQ2 examines whether cooperative post-processing meaningfully changes the behaviour of the Raw LLM controller.",
    )
    set_paragraph_text(
        doc.paragraphs[284],
        "The corrected evidence does not show a clear traffic-performance advantage for Hybrid over Raw LLM.",
    )
    set_paragraph_text(
        doc.paragraphs[285],
        "The 4V and 8V valid cells are effectively similar on the traffic metrics, and the corrected evidence contains no visible postprocessor intervention.",
    )
    set_paragraph_text(
        doc.paragraphs[286],
        "The cooperative stage is therefore present in the architecture but not empirically distinguished in the retained runs.",
    )
    set_paragraph_text(
        doc.paragraphs[287],
        "The correct reading is that the stage is implemented and traceable, but its effect is not measurable in the final evidence set.",
    )
    set_paragraph_text(
        doc.paragraphs[288],
        "This leaves no basis for claiming that cooperative processing alone improved traffic outcomes in the dissertation experiments.",
    )
    set_paragraph_text(
        doc.paragraphs[289],
        "Provider reliability also remains weak for the hybrid controllers, so the lack of separation is consistent with fallback dominance.",
    )
    set_paragraph_text(doc.paragraphs[290], "6.3 Safety Verification and RQ3")
    set_paragraph_text(
        doc.paragraphs[291],
        "RQ3 considers the effect of deterministic safety verification.",
    )
    set_paragraph_text(
        doc.paragraphs[292],
        "No collisions occurred in the retained experiments, but no safety overrides were recorded either.",
    )
    set_paragraph_text(
        doc.paragraphs[293],
        "The safety verifier was therefore implemented as part of the architecture but was not sufficiently exercised by the tested scenarios to establish either a safety benefit or a safety-efficiency trade-off.",
    )
    set_paragraph_text(
        doc.paragraphs[294],
        "The result is still useful because it shows that the safety stage can be traced, but the evidence boundary does not support a stronger claim.",
    )
    set_paragraph_text(
        doc.paragraphs[295],
        "In particular, the absence of collisions is an observation about the tested scenarios, not proof that the safety verifier improved safety.",
    )
    set_paragraph_text(
        doc.paragraphs[296],
        "A defensible formulation is that the safety layer was operationally present but not meaningfully triggered in the corrected formal evidence.",
    )
    set_paragraph_text(
        doc.paragraphs[297],
        "That formulation keeps the discussion aligned with the trace data and avoids overstating the role of the verifier.",
    )
    set_paragraph_text(
        doc.paragraphs[298],
        "The main safety conclusion is therefore negative in an evidential sense: the layer exists, but the dataset does not exercise it enough to measure its independent effect.",
    )
    set_paragraph_text(
        doc.paragraphs[299],
        "This is consistent with the low-density scenarios used in the formal experiment.",
    )
    set_paragraph_text(doc.paragraphs[300], "6.4 Vehicle Scale and RQ4")
    set_paragraph_text(
        doc.paragraphs[301],
        "RQ4 examines behaviour as the scenario increases from four to eight vehicles.",
    )
    set_paragraph_text(
        doc.paragraphs[302],
        "Traffic-level behaviour remains comparatively stable for the LLM-assisted pipelines, while the rule-based baseline degrades substantially at the larger scale.",
    )
    set_paragraph_text(
        doc.paragraphs[303],
        "Rule-based performance degrades substantially from 4V to 8V, whereas the LLM-assisted traffic metrics remain comparatively stable across the tested range.",
    )
    set_paragraph_text(
        doc.paragraphs[304],
        "Live-provider reliability remains weak at both scales, and the corrected 8V evidence is especially fallback-heavy.",
    )
    set_paragraph_text(
        doc.paragraphs[305],
        "That means the apparent scale stability belongs to the complete pipeline rather than to live model output alone.",
    )
    set_paragraph_text(
        doc.paragraphs[306],
        "The bounded dissertation claim is therefore limited to the tested 4V and 8V range.",
    )
    set_paragraph_text(
        doc.paragraphs[307],
        "Within that range, the LLM-assisted pipeline remained operational and retained strong traffic metrics.",
    )
    set_paragraph_text(
        doc.paragraphs[308],
        "The study does not show general scalability to denser traffic, larger vehicle populations, or more complex road networks.",
    )
    set_paragraph_text(doc.paragraphs[309], "6.5 Provider Reliability and Attribution")
    set_paragraph_text(
        doc.paragraphs[310],
        "Provider reliability is the main interpretive limitation of the dissertation. In the corrected 8V experiments, provider attempts total 2784 and provider successes total 4, corresponding to a success rate of approximately 0.14% and a fallback rate of approximately 99.86%.",
    )
    set_paragraph_text(
        doc.paragraphs[311],
        "Provider failures dominate the live history, which makes fallback-heavy execution the norm rather than the exception.",
    )
    set_paragraph_text(
        doc.paragraphs[312],
        "The surviving live-provider calls are therefore too sparse to support strong claims about model-only control quality.",
    )
    set_paragraph_text(
        doc.paragraphs[313],
        "What the dissertation does show is that the pipeline can remain operational despite severe provider unreliability.",
    )
    set_paragraph_text(
        doc.paragraphs[314],
        "The appropriate attribution is consequently pipeline-level rather than LLM-only.",
    )
    set_paragraph_text(
        doc.paragraphs[315],
        "This framing is necessary to keep the discussion consistent with the trace evidence and with the frozen evaluation design.",
    )
    set_paragraph_text(doc.paragraphs[316], "6.6 Relationship to Existing Literature")
    set_paragraph_text(
        doc.paragraphs[317],
        "The dissertation aligns with the recovered literature that treats autonomous intersection management as a structured decision problem and with recent work on staged LLM reasoning for embodied systems.",
    )
    set_paragraph_text(
        doc.paragraphs[318],
        "In particular, the results support the argument that language models are more defensible as one component in a traceable decision pipeline than as unconstrained controllers.",
    )
    set_paragraph_text(
        doc.paragraphs[319],
        "The corrected evidence also reinforces earlier findings that system behaviour depends strongly on the surrounding control logic, access constraints, and validation stages rather than on language generation alone.",
    )
    set_paragraph_text(
        doc.paragraphs[320],
        "That makes the dissertation consistent with the literature on modular reasoning, validation, and simulation-based traffic evaluation.",
    )
    set_paragraph_text(
        doc.paragraphs[321],
        "It does not provide a new claim of pure LLM superiority, which would exceed the evidence available here.",
    )
    set_paragraph_text(
        doc.paragraphs[322],
        "Instead, it adds a reproducible case in which provider availability and fallback control are explicitly visible in the analysis.",
    )
    set_paragraph_text(
        doc.paragraphs[323],
        "That is the contribution that most clearly connects the recovered literature to the dissertation's own experimental design.",
    )
    set_paragraph_text(doc.paragraphs[324], "6.7 Overall Interpretation")
    set_paragraph_text(
        doc.paragraphs[325],
        "The overall interpretation is that the dissertation demonstrates a functioning, traceable LLM-assisted decision pipeline whose traffic behaviour is better than the rule-based baseline in the tested scenarios, but whose live-provider layer is too unreliable to support any strong model-only claim.",
    )
    set_paragraph_text(
        doc.paragraphs[326],
        "The retained evidence therefore supports a pipeline-level reading: the architecture is useful, the traffic results are favourable, and the reliability constraints are central to understanding why the final controller behaviour looks the way it does.",
    )
    set_paragraph_text(
        doc.paragraphs[327],
        "The same evidence also rejects several stronger claims.",
    )
    set_paragraph_text(
        doc.paragraphs[328],
        "It does not show intrinsic LLM superiority, safety superiority, measurable cooperative-postprocessing benefit, or broad scalability beyond the tested low-density scenarios.",
    )
    set_paragraph_text(
        doc.paragraphs[329],
        "The most defensible dissertation conclusion is therefore bounded but clear.",
    )
    set_paragraph_text(
        doc.paragraphs[330],
        "A structured LLM-assisted pipeline can be evaluated systematically in SUMO, and in the tested scenarios it can outperform the rule-based baseline on the traffic metrics while remaining constrained by provider unreliability.",
    )
    set_paragraph_text(
        doc.paragraphs[331],
        "That conclusion is consistent with the results, the limitations, and the recovered literature.",
    )
    set_paragraph_text(
        doc.paragraphs[332],
        "It is also the strongest claim supported by the final evidence boundary.",
    )
    set_paragraph_text(
        doc.paragraphs[333],
        "The discussion therefore closes at the pipeline level rather than at the level of a standalone language model.",
    )
    set_paragraph_text(
        doc.paragraphs[334],
        "RQ1 is supported cautiously at the pipeline level.",
    )
    set_paragraph_text(
        doc.paragraphs[335],
        "RQ2 shows no clear traffic-performance advantage for Hybrid over Raw LLM.",
    )
    set_paragraph_text(
        doc.paragraphs[336],
        "RQ3 shows that the safety layer was present but insufficiently exercised.",
    )
    set_paragraph_text(
        doc.paragraphs[337],
        "RQ4 shows traffic robustness from 4V to 8V, but not broader scalability beyond the tested range.",
    )
    set_paragraph_text(doc.paragraphs[338], " ")

    # Chapter 7 / 8 light cleanup
    set_paragraph_text(
        doc.paragraphs[347],
        "In the corrected 8V experiments, provider success was extremely low: only 4 of 2,784 live-provider requests succeeded.",
    )
    set_paragraph_text(
        doc.paragraphs[361],
        "The corrected 8V analysis uses the independently executed final 8V evidence set rather than the invalid nominal 8V traces.",
    )
    clear_paragraph(doc.paragraphs[365])

    # References cleanup
    set_paragraph_text(
        doc.paragraphs[388],
        "Cui, C., Ma, Y., Yang, Z., Zhou, Y., Liu, P., Lu, J., Li, L., Chen, Y., Panchal, J.H., Abdelraouf, A., Gupta, R., Han, K. and Wang, Z. (2025) 'Large language models for autonomous driving (LLM4AD): Concept, benchmark, experiments, and challenges', arXiv preprint, arXiv:2410.15281v3.",
    )
    set_paragraph_text(
        doc.paragraphs[389],
        "Dong, X., Li, J., Xie, J., Yi, Y., Jia, T., Fang, S., Tian, Y. and Hang, P. (2026) 'Large language model based interactive decision-making for autonomous driving', arXiv preprint, arXiv:2604.23513v1.",
    )
    set_paragraph_text(
        doc.paragraphs[390],
        "Dresner, K. and Stone, P. (2008) 'A multiagent approach to autonomous intersection management', Journal of Artificial Intelligence Research, 31, pp. 591-656.",
    )
    set_paragraph_text(
        doc.paragraphs[391],
        "Driess, D., Xia, F., Sajjadi, M.S.M., Lynch, C., Chowdhery, A., Ichter, B., Wahid, A., Tompson, J., Vuong, Q., Yu, T., Huang, W., Chebotar, Y., Sermanet, P., Duckworth, D., Levine, S., Vanhoucke, V., Hausman, K., Toussaint, M., Greff, K., Zeng, A., Mordatch, I. and Florence, P. (2023) 'PaLM-E: An embodied multimodal language model', arXiv preprint, arXiv:2303.03378v1.",
    )
    set_paragraph_text(
        doc.paragraphs[392],
        "Hou, X., Wang, W., Yang, L., Lin, H., Feng, J., Min, H. and Zhao, X. (2025) 'DriveAgent: Multi-agent structured reasoning with LLM and multimodal sensor fusion for autonomous driving', arXiv preprint, arXiv:2505.02123v1.",
    )
    set_paragraph_text(
        doc.paragraphs[393],
        "Huang, W., Abbeel, P., Pathak, D. and Mordatch, I. (2022) 'Language models as zero-shot planners: Extracting actionable knowledge for embodied agents', arXiv preprint, arXiv:2201.07207v2.",
    )
    set_paragraph_text(
        doc.paragraphs[394],
        "Safarov, K. (2022) The impact of autonomous vehicles on traffic performance at an unregulated junction. PhD thesis. University of Bristol.",
    )

    # Appendices
    set_paragraph_text(doc.paragraphs[407], "Appendix C. Experimental Matrix and Seed-level Results")
    set_paragraph_text(
        doc.paragraphs[408],
        "This appendix preserves the final evidence boundary and the seed-level values used to assemble the dissertation tables.",
    )
    set_paragraph_text(doc.paragraphs[409], "Table C1. Seed-level traffic values used to assemble Table 2")
    set_paragraph_text(doc.paragraphs[413], "Appendix E. Abbreviations")
    set_paragraph_text(doc.paragraphs[414], "Table E1. Abbreviations used in the dissertation")

    # Table 2: simplify and remove seed-level values column
    table2 = doc.tables[2]
    remove_last_column(table2)
    headers = [
        "Controller",
        "Scale",
        "Completion rate",
        "Mean waiting time (mean ± SD)",
        "Mean speed (mean ± SD)",
        "Throughput",
        "Collisions",
    ]
    for c, text in enumerate(headers):
        set_table_cell(table2, 0, c, text)
    table2_rows = [
        ["Rule-based", "4V", "100%", "82.000 ± 0.000 steps", "2.310 ± 0.000 m/s", "4.000 ± 0.000", "0"],
        ["Rule-based", "8V", "100%", "242.042 ± 110.586 steps", "1.189 ± 0.754 m/s", "8.000 ± 0.000", "0"],
        ["Raw LLM", "4V", "100%", "15.000 ± 0.000 steps", "6.803 ± 0.000 m/s", "4.000 ± 0.000", "0"],
        ["Raw LLM", "8V", "100%", "15.292 ± 2.045 steps", "6.599 ± 0.254 m/s", "8.000 ± 0.000", "0"],
        ["Hybrid", "4V", "100%", "15.000 ± 0.000 steps", "6.803 ± 0.000 m/s", "4.000 ± 0.000", "0"],
        ["Hybrid", "8V", "100%", "15.292 ± 2.045 steps", "6.599 ± 0.254 m/s", "8.000 ± 0.000", "0"],
        ["Hybrid + Safety", "4V", "100%", "15.000 ± 0.000 steps", "6.803 ± 0.000 m/s", "4.000 ± 0.000", "0"],
        ["Hybrid + Safety", "8V", "100%", "15.292 ± 2.045 steps", "6.599 ± 0.254 m/s", "8.000 ± 0.000", "0"],
    ]
    for r, row in enumerate(table2_rows, start=1):
        for c, text in enumerate(row):
            set_table_cell(table2, r, c, text)

    # Appendix C table becomes seed-level raw values
    table7 = doc.tables[7]
    table7.rows[0].cells[0].text = "Controller / Scale"
    table7.rows[0].cells[1].text = "Seed-level raw values"
    seed_rows = [
        ["Rule-based 4V", "completion [1.0, 1.0, 1.0]; waiting [82, 82, 82]; speed [2.310, 2.310, 2.310]"],
        ["Rule-based 8V", "completion [1.0, 1.0, 1.0]; waiting [86, 311, 329.125]; speed [2.255, 0.655, 0.658]"],
        ["Raw LLM 4V", "completion [1.0, 1.0, 1.0]; waiting [15, 15, 15]; speed [6.803, 6.803, 6.803]"],
        ["Raw LLM 8V", "completion [1.0, 1.0, 1.0]; waiting [17.875, 12.875, 15.125]; speed [6.265, 6.880, 6.652]"],
        ["Hybrid 4V", "completion [1.0, 1.0, 1.0]; waiting [15, 15, 15]; speed [6.803, 6.803, 6.803]"],
        ["Hybrid 8V", "completion [1.0, 1.0, 1.0]; waiting [17.875, 12.875, 15.125]; speed [6.265, 6.880, 6.652]"],
        ["Hybrid + Safety 4V", "completion [1.0, 1.0, 1.0]; waiting [15, 15, 15]; speed [6.803, 6.803, 6.803]"],
        ["Hybrid + Safety 8V", "completion [1.0, 1.0, 1.0]; waiting [17.875, 12.875, 15.125]; speed [6.265, 6.880, 6.652]"],
    ]
    add_rows(table7, 9)
    for r, row in enumerate(seed_rows, start=1):
        table7.rows[r].cells[0].text = row[0]
        table7.rows[r].cells[1].text = row[1]

    # Appendix E abbreviations
    table9 = doc.tables[9]
    remove_last_rows(table9, 2)
    abbreviations = [
        ["LLM", "Large language model"],
        ["SUMO", "Simulation of Urban MObility"],
        ["RQ", "Research question"],
    ]
    for r, row in enumerate(abbreviations, start=1):
        table9.rows[r].cells[0].text = row[0]
        table9.rows[r].cells[1].text = row[1]

    set_update_fields(doc)
    doc.save(str(DOCX_TMP))

    # Replace Figure 3 image payload while keeping the existing relationship.
    shutil.copyfile(DOCX_TMP, DOCX_OUT)
    patch_media(DOCX_OUT, "word/media/image3.png", FIG3_OUT)
    if DOCX_TMP.exists():
        DOCX_TMP.unlink()


if __name__ == "__main__":
    main()
